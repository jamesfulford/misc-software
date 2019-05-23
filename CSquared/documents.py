import os
import shutil

import docx
from jinja2 import Template

from document_contexts import *
# from writers import *

#
# Constants
#


def replace_variables(s, mapping, prefmter=None):
    """
    Used to replace variables with values
    """
    if prefmter is None:
        prefmter = lambda s: s

    replacings = []
    replacings = map(lambda i: (prefmter(i[0]), str(i[1])), mapping.items())
    # replace larger before smaller, to prevent substring problems
    replacings = sorted(replacings, key=lambda i: -len(i[0]))

    st = s
    for before, after in replacings:
        st = st.replace(before, after)
    return st


def docx_replace_variables(doc, mapping):
    """
    Replaces all instances of key in document with corresponding value
    """
    for p in doc.paragraphs:
        for run in p.runs:
            run.text = replace_variables(
                run.text, mapping, prefmter=lambda s: s
            )

    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_variables(cell, mapping)


def transform_document(old_file, new_file, context={}, jinja=True):
    if ".docx" in old_file:
        doc = docx.Document(old_file)
        docx_replace_variables(doc, context)
        doc.save(new_file)
    elif ".txt" in old_file:
        r = ""
        with open(old_file, "r") as read:
            r = read.read()

        if jinja:
            r = Template(r).render(context)
        else:
            r = replace_variables(r, context)

        with open(new_file, "w") as write:
            write.write(r)
    else:
        shutil.copyfile(old_file, new_file)

    return new_file


def just_copy(old_file, new_file, context=None):
    """
    Use to exempt a file from being transformed.
    """
    shutil.copyfile(old_file, new_file)


def rewrite_with(context_maker):
    """
    Convenient tool for hooking context_maker functions into PATTERN_REPORTMAKER
    Context_getter returns context given post
    """
    return lambda o, n, p: transform_document(
        o, n, context=context_maker(p), jinja=False
    )


# matches OLD_FILE (before transforming based on file_context variables)
PATTERN_REPORTMAKER = {
    "Post Report.docx": just_copy,
    "NOC Posting Procedure_rev5(Macro).docx": rewrite_with(
        posting_procedure_context
    ),
    "TEMPLATE Post Report Template.docx": just_copy,
    # "One Time Requests.docx": write_one_time_requests
}


def transform_directory(dirpath, target, post, IGNORE=[]):
    # TODO: regex match instead of substring check
    files = filter(
        lambda i: i not in IGNORE,
        os.listdir(dirpath)
    )
    directories = filter(
        lambda f: os.path.isdir(os.path.join(dirpath, f)),
        files
    )
    files = filter(
        lambda f: not os.path.isdir(os.path.join(dirpath, f)),
        files
    )

    try:
        os.mkdir(target)
    except OSError:
        pass

    for f in files:
        old_file = os.path.join(dirpath, f)
        new_file = os.path.join(
            target,
            replace_variables(
                f, file_renaming_context(post)
            )
        )

        # TODO: regex match instead of substring check
        matched = False
        for pattern, rewriter in PATTERN_REPORTMAKER.items():
            if pattern in old_file:
                rewriter(old_file, new_file, post)
                matched = True
                break
        if not matched:
            transform_document(
                old_file, new_file, context=generate_document_context(post)
            )

    for d in directories:
        transform_directory(
            os.path.join(dirpath, d),
            os.path.join(target, d),
            post,
            IGNORE=IGNORE
        )


if __name__ == "__main__":
    print replace_variables("$target_environment", {"$target": "shofar"})
