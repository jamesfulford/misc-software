# sandbox.py
# by James Fulford

import analytics as a
import json
import matplotlib.pyplot as plt
import docx
from datetime import datetime as dt


def frequency(data, count):
    # Making a frequency table by classes:
    class_size = a.datarange(data) / float(count)

    # Generating the classes:
    classes = []
    bottom = a.minimum(data)
    for cl in range(count):
        classes.append((bottom, round(bottom + class_size, 10)))
        bottom = round(bottom + class_size, 10)

    def str_class(clas):
        # represent which class as a string...
        key = "[" + str(clas[0]) + ", " + str(clas[1])
        # represent differently if it's the last class:
        if classes.index(clas) == len(classes) - 1:
            # last class is closed interval
            key += "]"
        else:
            key += ")"
        return key

    # Counting number of items in each class
    freq = {}  # holds classes as keys, count as values
    for cl in classes:
        freq[str_class(cl)] = 0  # Set all counts to 0's

    # We will count every dat in the data.
    for dat in data:
        # Which class does this dat fall into?
        # Let's go through all of them.
        for i in range(len(classes)):
            cl = classes[i]
            # Let's check if dat is inside this class:
            # if lower <= dat < higher, it's in the class:
            if dat >= cl[0] and dat < cl[1]:
                freq[str_class(cl)] += 1  # add one to the class
                # and move onto the next piece of data
                break

            # if dat is not in any class
            # but close to upper class:
            if i == len(classes) - 1 and cl[1] - dat < 0.0001:
                freq[str_class(cl)] += 1  # add 1 to final class
    return freq


# ready-up a Word document:
doc = docx.Document()
doc.add_heading("Statistics", 0)
doc.add_heading("By James Fulford\tProblem Solving and Modelling\t" + dt.now().strftime("%m/%d/%Y"))
doc.add_heading("Data", 1)

# Cereal Box Weight Data
data = [16.22, 15.79, 16.03, 15.76, 15.72, 15.89, 16.25, 15.88, 15.89, 15.62]
data.extend([15.95, 15.89, 15.99, 16.03, 15.94, 16.29, 16.06, 16.17, 15.97, 16.16])
data.extend([16.14, 16.00, 16.16, 15.96, 15.89, 16.08, 15.94, 15.96, 16.08, 16.06])
data.extend([16.17, 16.18, 16.10, 15.98, 16.17, 15.93, 15.87, 15.92, 16.03, 16.05])
data.extend([15.98, 15.99, 15.92, 16.22, 16.25, 15.64, 16.03, 16.05, 16.26, 16.14])
data.sort()
doc.add_paragraph("Cereal Box Weights: " + json.dumps(data))

# X-Files Viewership Data
xdata = [9.8, 9.3, 8.7, 8.2, 8.5, 9.2, 9, 9.1, 8.6, 9.9]
xdata.extend([8.5, 9.7, 8.8, 10.2, 10.8, 9.8, 10.7, 9.6, 10.2, 9.8])
xdata.extend([7.9, 8.5, 8.1, 9, 9.6])
xdata.sort()
doc.add_paragraph("X-Files Viewership: " + json.dumps(xdata))


##############################
##############################
##############################
# Next problem... skip two lines.
print "\n" * 1
doc.add_page_break()
##############################

# Question 3 b)

count = 8
data_frequency = frequency(data, count)

print "3. b) Frequency Table (classes: " + str(count) + ")"
print(json.dumps(data_frequency, indent=4, sort_keys=lambda x: x))  # this shows the frequency table

# MAKE HISTOGRAM
print "Showing histogram now."
plt.hist(data, count)  # this does most of the work above for us.
plt.axis([a.minimum(data), a.maximum(data), 0, a.maximum(data_frequency.values()) * 1.1])

# Labels and Titles
plt.title("Cereal Box Weights Histogram (classes: " + str(count) + ")")
plt.xlabel("Weight (ounces)")
plt.ylabel("Frequency (n: " + str(a.count(data)) + ")")

# Show to user
plt.savefig("Cereal Box Weights Histogram.png")
plt.show()


# Write up the document:
doc.add_heading("Problem 3 b)", 1)
doc.add_paragraph("Frequency Table (classes: " + str(count) + ")")
doc.add_paragraph(json.dumps(data_frequency, indent=4, sort_keys=lambda x: x))
doc.add_paragraph("Cereal Box Weights Histogram:")
doc.add_picture("Cereal Box Weights Histogram.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))


##############################
##############################
##############################
# Next problem... skip two lines.
print "\n" * 1
doc.add_page_break()
##############################

# Question 3 c)
count = 8
xdata_frequency = frequency(xdata, count)

print "3. c) X-Files Frequency Table (classes: " + str(count) + ")"
print(json.dumps(xdata_frequency, indent=4, sort_keys=lambda x: x))  # this shows the frequency table

# MAKE HISTOGRAM
print "Showing X-Files histogram now."
plt.hist(xdata, count)  # this does most of the work above for us.
plt.axis([a.minimum(xdata), a.maximum(xdata), 0, a.maximum(xdata_frequency.values()) * 1.1])

# Labels and Titles
plt.title("X-Files Viewers (Second Season) Histogram (classes: " + str(count) + ")")
plt.xlabel("Viewership (millions)")
plt.ylabel("Frequency (n: " + str(a.count(xdata)) + ")")

# Show to user
plt.savefig("X-Files Viewers (Second Season) Histogram.png")
plt.show()


# Write up the document:
doc.add_heading("Problem 3 c)", 1)
doc.add_paragraph("Frequency Table (classes: " + str(count) + ")")
doc.add_paragraph(json.dumps(data_frequency, indent=4, sort_keys=lambda x: x))
doc.add_paragraph("Histogram of X-Files Viewership:")
doc.add_picture("X-Files Viewers (Second Season) Histogram.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))


##############################
##############################
##############################
# Next problem... skip two lines.
print "\n" * 1
doc.add_page_break()
##############################


# Show Five-point Summary
print "4. Five-point Summary:"
print a.five_point_summary(xdata)
print "Showing boxplot now."
print

# Make a boxplot of the data
plt.boxplot(xdata)

# Labels and Titles
plt.title("X-Files Viewers (Second Season)")
plt.ylabel("Viewers (millions)")
plt.xlabel("")

# Show to user
plt.savefig("X-Files Viewers (Second Season) Boxplot.png")
plt.show()

# Write to document
doc.add_heading("Problem 4", 1)
doc.add_paragraph("Five Point Summary: " + str(a.five_point_summary(xdata)))
doc.add_paragraph("Boxplot of X-Files Viewers (Second Season):")
doc.add_picture("X-Files Viewers (Second Season) Boxplot.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))

##############################
##############################
##############################

# Save document
doc.save("Statistics " + dt.now().strftime("%m-%d-%Y") + " James Fulford.docx")
